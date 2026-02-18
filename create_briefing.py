import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL

# --- CONFIGURATION ---
NOM_FICHIER_SORTIE = "Dossier_Vol_IFR_CPL.docx"

# Images à insérer (Doivent être présentes dans le dossier)
IMAGES_METEO = [
    ("SITUATION GENERALE (WAFC)", "fronts_wafc.png"),
    ("TEMSI FRANCE", "temsi_france.png"),
    ("WINTEM (Vents/Temp)", "wintem_france.png"),
    ("SIGMET (Sécurité)", "sigmet_france.png")
]

def set_cell_shading(cell, color_hex):
    """Applique un fond gris ou couleur à une cellule"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_header_cell(cell, text, bold=True):
    """Formate une cellule d'en-tête (Gris + Gras)"""
    set_cell_shading(cell, "D9D9D9") # Gris standard Word
    p = cell.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

def add_section_title(doc, text):
    """Ajoute un titre de section propre (Style Identique Doc Original)"""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(2)
    runner = p.add_run(text)
    runner.bold = True
    runner.font.size = Pt(12)
    runner.font.name = 'Arial'
    runner.font.color.rgb = RGBColor(0, 0, 0) # Noir
    # Ligne de séparation
    p_border = doc.add_paragraph()
    p_border.paragraph_format.line_spacing = Pt(0)
    p_border.paragraph_format.space_after = Pt(6)
    run_border = p_border.add_run("_________________________________________________________________________________")
    run_border.font.size = Pt(8)
    run_border.font.color.rgb = RGBColor(150, 150, 150)

def create_document():
    doc = Document()
    
    # 1. MISE EN PAGE (Marges optimisées)
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)

    # --- EN-TÊTE (Header Logique) ---
    table_head = doc.add_table(rows=2, cols=2)
    table_head.autofit = False
    table_head.columns[0].width = Cm(14)
    table_head.columns[1].width = Cm(5)
    
    # Titre Gauche
    c1 = table_head.cell(0, 0)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run("FLIGHT BRIEFING IFR / VFR")
    r1.bold = True
    r1.font.size = Pt(16)
    
    # Cases à cocher Droite
    c2 = table_head.cell(0, 1)
    p2 = c2.paragraphs[0]
    p2.add_run("DATE : ______________\n")
    p2.add_run("☐ NCO (VFR)   ☐ CAT (IFR)")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Ligne Pilotes / Avion
    row2 = table_head.rows[1]
    p_pilote = row2.cells[0].paragraphs[0]
    p_pilote.add_run("AVION : ___________   CDB : ___________________   INSTR : ___________________")
    
    doc.add_paragraph("")

    # --- 1. FLIGHT DOCUMENTS (Identique Original) ---
    add_section_title(doc, "1. FLIGHT DOCUMENTS & ADMIN")
    
    table_docs = doc.add_table(rows=1, cols=4)
    table_docs.style = 'Table Grid'
    
    # Colonne 1 : Documents Avion
    c = table_docs.cell(0, 0)
    c.text = "AIRCRAFT DOCUMENTS"
    set_cell_shading(c, "E7E6E6") # Gris très clair
    p = c.paragraphs[0]
    p.runs[0].font.bold = True
    
    content_docs = [
        "☐ Insurance", "☐ Registration", "☐ Noise Cert.", 
        "☐ Airworthiness (CDN)", "☐ ARC / CEN", "☐ Radio Licence (LSA)",
        "☐ Tech Log (Potentiel OK)", "☐ Mass & Balance Sheet", "☐ AFM (Manual)"
    ]
    c_list = table_docs.cell(0, 1)
    for item in content_docs:
        c_list.add_paragraph(item).paragraph_format.space_after = Pt(0)

    # Colonne 2 : Crew & Safety
    c_crew = table_docs.cell(0, 2)
    c_crew.text = "CREW & SAFETY"
    set_cell_shading(c_crew, "E7E6E6")
    c_crew.paragraphs[0].runs[0].font.bold = True

    content_safety = [
        "☐ Licences & Medical", "☐ Maps / Charts (Paper/EFB)", 
        "☐ Flight Plan (Déposé)", "☐ Gilet / Life Vest", "☐ Trousse Secours"
    ]
    c_safe = table_docs.cell(0, 3)
    for item in content_safety:
        c_safe.add_paragraph(item).paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph("")

    # --- 2. DEPARTURE (Chronologie Instructeur) ---
    add_section_title(doc, "2. DEPARTURE")
    
    t_dep = doc.add_table(rows=2, cols=2)
    t_dep.style = 'Table Grid'
    
    # Ligne 1 : Terrains
    t_dep.cell(0,0).text = "TERRAIN DÉPART (ICAO) : ________"
    t_dep.cell(0,1).text = "DÉGAGEMENT (ALTN DEP) : ________"
    
    # Ligne 2 : NOTAM & METEO (Briefing Point)
    c_notam = t_dep.cell(1,0)
    c_notam.text = "NOTAM & PISTE EN SERVICE :\n\n\n"
    
    c_meteo = t_dep.cell(1,1)
    c_meteo.text = "METEO DÉPART (METAR / TAF) :\n\n\n"
    
    doc.add_paragraph("")

    # --- 3. EN ROUTE (Avec le tableau demandé "EN PLUS") ---
    add_section_title(doc, "3. EN ROUTE")
    
    # Tableau 1 : Info Route
    t_route_info = doc.add_table(rows=1, cols=1)
    t_route_info.style = 'Table Grid'
    t_route_info.cell(0,0).text = "NOTAM EN ROUTE (ZRT, AZBA, NAVAIDS...) & PARTICULARITÉS :\n\n\n"
    
    doc.add_paragraph("")
    
    # Tableau 2 : LE TABLEAU "FUEL LOG" DEMANDÉ PAR L'INSTRUCTEUR
    p_log = doc.add_paragraph()
    p_log.add_run("NAVIGATION LOG & SUIVI CARBURANT :").bold = True
    p_log.paragraph_format.space_after = Pt(2)
    
    t_log = doc.add_table(rows=6, cols=5)
    t_log.style = 'Table Grid'
    
    # Headers
    headers_log = ["POINT / WAYPOINT", "ROUTE / AWY", "FL / ALT", "ETO (Temps)", "FUEL ESTIMÉ"]
    for i, h in enumerate(headers_log):
        format_header_cell(t_log.cell(0, i), h)
        
    # Lignes vides pour remplissage
    for r in range(1, 6):
        for c in range(5):
            t_log.cell(r, c).text = ""

    doc.add_paragraph("")

    # --- 4. DESTINATION ---
    add_section_title(doc, "4. DESTINATION & ARRIVAL")
    
    t_dest = doc.add_table(rows=2, cols=2)
    t_dest.style = 'Table Grid'
    t_dest.cell(0,0).text = "TERRAIN DESTINATION : ________"
    t_dest.cell(0,1).text = "DÉGAGEMENT (ALTN) : ________"
    
    t_dest.cell(1,0).text = "NOTAM ARRIVÉE & APPROCHE PRÉVUE :\n\n\n"
    t_dest.cell(1,1).text = "METEO ARRIVÉE (METAR / TAF) :\n\n\n"
    
    doc.add_paragraph("")
    doc.add_page_break() # Nouvelle page pour la météo et la suite technique

    # --- 5. MÉTÉO GLOBALE (Le Robot) ---
    add_section_title(doc, "5. SITUATION MÉTÉO GÉNÉRALE (CARTES ROBOT)")
    
    # Tableau 2x2 pour les cartes
    t_img = doc.add_table(rows=2, cols=2)
    t_img.autofit = True
    
    for i, (titre, f_img) in enumerate(IMAGES_METEO):
        r, c = i // 2, i % 2
        cell = t_img.cell(r, c)
        
        # Titre Carte
        p = cell.paragraphs[0]
        p.add_run(titre).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Insertion Image
        if os.path.exists(f_img):
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            try:
                run.add_picture(f_img, width=Cm(8.0))
            except:
                p_img.add_run(f"[Erreur format {f_img}]")
        else:
            cell.add_paragraph(f"\n[IMAGE {f_img} ABSENTE]\nLancer le robot météo avant.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # --- 6. PLANNING MINIMA (IFR) ---
    add_section_title(doc, "6. PLANNING MINIMA")
    
    t_min = doc.add_table(rows=3, cols=5)
    t_min.style = 'Table Grid'
    
    headers_min = ["TERRAIN", "TYPE APPROCHE", "DA / MDA", "VISIBILITÉ", "PLAFOND"]
    for i, h in enumerate(headers_min):
        format_header_cell(t_min.cell(0, i), h)
        
    row_labels = ["DESTINATION", "ALTERNATE"]
    for i, label in enumerate(row_labels):
        c = t_min.cell(i+1, 0)
        c.text = label
        c.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(c, "F2F2F2")

    doc.add_paragraph("")

    # --- 7. MASSE & CENTRAGE (Déplacé AVANT le Fuel - Demande Instructeur) ---
    add_section_title(doc, "7. MASSE & CENTRAGE (M&B)")
    
    t_mb = doc.add_table(rows=1, cols=4)
    t_mb.style = 'Table Grid'
    
    t_mb.cell(0,0).text = "MASSE À VIDE :\n________ kg"
    t_mb.cell(0,1).text = "CHARGE UTILE :\n________ kg"
    t_mb.cell(0,2).text = "MASSE DÉCOLLAGE (TOW) :\n________ kg (Max: ____)"
    t_mb.cell(0,3).text = "CENTRAGE DÉPART :\n☐ OK    ☐ HORS LIMITES"
    
    doc.add_paragraph("Nota : Vérifier la masse atterrissage si vol court.")
    doc.add_paragraph("")

    # --- 8. CARBURANT (FUEL) ---
    add_section_title(doc, "8. CARBURANT")
    
    t_fuel = doc.add_table(rows=7, cols=4)
    t_fuel.style = 'Table Grid'
    
    # Headers
    format_header_cell(t_fuel.cell(0, 0), "ITEM")
    format_header_cell(t_fuel.cell(0, 1), "TEMPS")
    format_header_cell(t_fuel.cell(0, 2), "QUANTITÉ (L)")
    format_header_cell(t_fuel.cell(0, 3), "QUANTITÉ (Gal)")
    
    fuel_items = [
        "TAXI (Roulage)",
        "TRIP (Départ -> Dest)",
        "CONTINGENCY (5% du Trip)",
        "ALTERNATE (Dégagement)",
        "FINAL RESERVE (45 min)",
        "EXTRA / DISCRETION"
    ]
    
    for i, item in enumerate(fuel_items):
        t_fuel.cell(i+1, 0).text = item
        # Petit grisage pour la Final Reserve (Vital)
        if "FINAL" in item:
            set_cell_shading(t_fuel.cell(i+1, 0), "E2EFDA") # Vert très pâle alerte
            
    # Ligne TOTAL
    # (On peut l'ajouter à la main ou laisser le pilote faire le total)
    
    doc.add_paragraph("")

    # --- 9. PERFORMANCES & TEM ---
    add_section_title(doc, "9. PERFORMANCES & TEM")
    
    t_perf = doc.add_table(rows=2, cols=2)
    t_perf.style = 'Table Grid'
    
    t_perf.cell(0,0).text = "DÉCOLLAGE (TODR/TORA) :\nDist. Calculée : ______ m\nPiste Dispo : ______ m"
    t_perf.cell(0,1).text = "ATTERRISSAGE (LDR/LDA) :\nDist. Calculée : ______ m\nPiste Dispo : ______ m"
    t_perf.cell(1,0).merge(t_perf.cell(1,1))
    
    # Cellule TEM
    c_tem = t_perf.cell(1,0)
    p_tem = c_tem.paragraphs[0]
    p_tem.add_run("THREATS (Menaces) : ").bold = True
    p_tem.add_run("Météo - Terrain - Trafic - État Avion - Pression Tempo\n")
    p_tem.add_run("ERRORS (Erreurs) : ").bold = True
    p_tem.add_run("Oublis - Violations - Pilotage - Procédures\n\n")
    
    r_dec = p_tem.add_run("DECISION PILOTE :   ☐ GO      ☐ DELAY      ☐ NO-GO")
    r_dec.bold = True
    r_dec.font.size = Pt(12)
    r_dec.font.color.rgb = RGBColor(0, 0, 0)
    p_tem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(c_tem, "FFF2CC") # Jaune pâle "Attention"

    # Sauvegarde
    try:
        doc.save(NOM_FICHIER_SORTIE)
        print(f"✅ SUCCÈS : {NOM_FICHIER_SORTIE} a été généré.")
        print(f"   Structure : Conforme Instructeur (Météo P5, M&B P7).")
        print(f"   Images    : Intégrées (Si robot lancé).")
    except PermissionError:
        print(f"❌ ERREUR : Fermez le fichier {NOM_FICHIER_SORTIE} avant de relancer le script !")

if __name__ == "__main__":
    create_document()
